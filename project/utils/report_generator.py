from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

def generate_pdf_report(report: dict, path: str):
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    c.setFont('Helvetica-Bold', 18)
    c.drawString(40, height - 60, 'PresentAI Pro - Analysis Report')
    c.setFont('Helvetica', 10)
    y = height - 100
    for k, v in report.items():
        text = f"{k}: {v}"
        c.drawString(40, y, text[:90])
        y -= 16
        if y < 80:
            c.showPage(); y = height - 60
    c.save()


def generate_docx_report(report: dict, path: str):
    doc = Document()
    doc.add_heading('PresentAI Pro - Analysis Report', level=1)
    for k, v in report.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.save(path)
